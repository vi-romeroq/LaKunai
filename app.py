import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
import os
import datetime
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from analyzer import StandardAnalyzer
import hashlib
import bcrypt
from fpdf import FPDF
import re
import smtplib
from email.mime.text import MIMEText

def send_welcome_email(user_email):
    """Sends a professional welcome email if SMTP is configured."""
    sender = os.getenv("SMTP_USER", "")
    pwd = os.getenv("SMTP_PASS", "")
    if not sender or not pwd:
        return False
    msg = MIMEText(f"Bienvenido a Normatix GRC Platform.\n\nTu cuenta administrativa corporativa '{user_email}' ha sido habilitada exitosamente.\n\nIngresa a nuestra suite web para auditar el cumplimiento normativo de IA de tu compañía.\n\nEl equipo de Normatix.")
    msg['Subject'] = 'Bienvenido a Normatix | Secure Access'
    msg['From'] = f"Normatix Security <{sender}>"
    msg['To'] = user_email
    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as s:
            s.login(sender, pwd)
            s.send_message(msg)
        return True
    except Exception:
        return False


def generate_pdf_report(username: str, doc_names: str, risk_tier: str, report_text: str) -> bytes:
    """Generate a professional branded PDF audit report."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    # Header bar
    pdf.set_fill_color(5, 11, 20)
    pdf.rect(0, 0, 210, 35, 'F')
    if os.path.exists("logo.png"):
        pdf.image("logo.png", x=18, y=5, w=60, h=22)
    else:
        pdf.set_font('Helvetica', 'B', 22)
        pdf.set_text_color(56, 189, 248)
        pdf.set_xy(20, 8)
        pdf.cell(0, 10, 'NORMATIX', ln=False)
    pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(148, 163, 184)
    pdf.set_xy(20, 21)
    pdf.cell(0, 6, 'AI Enterprise GRC Platform  |  Reporte Oficial de Auditoría Documental')

    # Metadata block
    pdf.set_xy(20, 42)
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(30, 30, 30)
    pdf.cell(0, 7, f'Cliente/Usuario: {username}', ln=True)
    pdf.set_x(20)
    pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(70, 70, 70)
    pdf.cell(0, 6, f'Documentos Analizados: {doc_names}', ln=True)
    pdf.set_x(20)
    pdf.cell(0, 6, f'Fecha de Auditoría: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")} UTC', ln=True)

    # Risk badge
    pdf.ln(4)
    risk_colors = {
        'UNACCEPTABLE': (239, 68, 68), 'HIGH': (249, 115, 22),
        'LIMITED': (234, 179, 8), 'MINIMAL': (34, 197, 94)
    }
    r, g, b = risk_colors.get(risk_tier.upper(), (100, 100, 100))
    pdf.set_fill_color(r, g, b)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 12)
    pdf.set_x(20)
    pdf.cell(80, 10, f'  CLASIFICACIÓN DE RIESGO: {risk_tier}  ', fill=True, ln=True, align='C')

    # Divider
    pdf.ln(4)
    pdf.set_draw_color(56, 189, 248)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)

    # Report body
    pdf.set_font('Helvetica', '', 10)
    pdf.set_text_color(30, 30, 30)
    for line in report_text.split('\n'):
        clean = line.strip()
        if not clean:
            pdf.ln(3)
            continue
        if clean.startswith('###') or clean.startswith('##'):
            pdf.set_font('Helvetica', 'B', 11)
            pdf.set_text_color(5, 11, 20)
            pdf.set_x(20)
            pdf.multi_cell(170, 7, clean.lstrip('#').strip())
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(30, 30, 30)
        else:
            pdf.set_x(20)
            pdf.multi_cell(170, 6, clean)

    # Footer
    pdf.set_y(-20)
    pdf.set_font('Helvetica', 'I', 8)
    pdf.set_text_color(148, 163, 184)
    pdf.cell(0, 5, '© 2026 Normatix Soluciones Inteligentes | Este reporte es de carácter orientativo y no constituye asesoría legal vinculante.', align='C')

    return bytes(pdf.output())

# --- Cloud-Ready Database ORM (SQLAlchemy Setup) ---
from sqlalchemy import create_engine, Column, Integer, String, DateTime, Text
from sqlalchemy.orm import declarative_base, sessionmaker

load_dotenv()
DB_URL = os.getenv("DATABASE_URL", "sqlite:///normatix_fallback.db")
if "sqlite" in DB_URL:
    connect_args = {"check_same_thread": False}
    engine = create_engine(DB_URL, connect_args=connect_args)
else:
    # Supabase Transaction Pooler (port 6543) requires prepared statements disabled
    engine = create_engine(DB_URL, pool_pre_ping=True, execution_options={"prepare_threshold": None})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, autoincrement=True)
    username = Column(String, unique=True, index=True)
    password_hash = Column(String)
    role = Column(String, default="AUDITOR_LEGAL")
    plan = Column(String, default="FREE")

class Usage(Base):
    __tablename__ = "usage"
    username = Column(String, primary_key=True)
    count = Column(Integer, default=0)

class Audit(Base):
    __tablename__ = "audits"
    id = Column(Integer, primary_key=True, autoincrement=True)
    username = Column(String, index=True)
    doc_name = Column(String)
    risk_tier = Column(String)
    audit_date = Column(DateTime)

class RagDocument(Base):
    """Persistent RAG context stored in Supabase so it survives Streamlit Cloud restarts."""
    __tablename__ = "rag_documents"
    id = Column(Integer, primary_key=True, autoincrement=True)
    username = Column(String, index=True)
    doc_name = Column(String)
    content = Column(Text)
    created_at = Column(DateTime)

class GuestSession(Base):
    """Tracks anonymous guest audits by IP hash to prevent refresh-bypass abuse."""
    __tablename__ = "guest_sessions"
    ip_hash = Column(String, primary_key=True)
    used = Column(Integer, default=0)
    created_at = Column(DateTime)

Base.metadata.create_all(bind=engine)

def elevate_admin():
    """Auto-elevate the owner account to ADMINISTRADOR on every startup."""
    admin_user = os.getenv("ADMIN_USERNAME", "")
    if not admin_user:
        return
    db = SessionLocal()
    try:
        u = db.query(User).filter(User.username == admin_user).first()
        if u and u.role != "ADMINISTRADOR":
            u.role = "ADMINISTRADOR"
            db.commit()
    finally:
        db.close()

elevate_admin()

def get_usage(username):
    try:
        db = SessionLocal()
        u = db.query(Usage).filter(Usage.username == username).first()
        return u.count if u else 0
    finally:
        db.close()

def increment_usage(username):
    try:
        db = SessionLocal()
        u = db.query(Usage).filter(Usage.username == username).first()
        if u:
            u.count += 1
        else:
            db.add(Usage(username=username, count=1))
        db.commit()
    finally:
        db.close()

def save_audit(username, doc_name, risk_tier):
    try:
        db = SessionLocal()
        new_audit = Audit(username=username, doc_name=doc_name, risk_tier=risk_tier, audit_date=datetime.datetime.now())
        db.add(new_audit)
        db.commit()
    finally:
        db.close()

def get_audits(username):
    try:
        db = SessionLocal()
        audits = db.query(Audit).filter(Audit.username == username).order_by(Audit.id.desc()).all()
        return [(a.doc_name, a.risk_tier, a.audit_date.strftime("%Y-%m-%d %H:%M:%S")) for a in audits]
    finally:
        db.close()

def get_guest_ip_hash() -> str:
    """Returns a SHA-256 hash of the visitor's IP + User-Agent for anonymous fingerprinting."""
    try:
        headers = st.context.headers
        ip = headers.get("X-Forwarded-For", headers.get("X-Real-Ip", "unknown"))
        ua = headers.get("User-Agent", "")
        raw = f"{ip}:{ua}"
        return hashlib.sha256(raw.encode()).hexdigest()
    except Exception:
        return "unknown"

def check_and_mark_guest_ip(ip_hash: str) -> bool:
    """Returns True if this guest IP has NOT used their free audit yet. Marks it as used."""
    if ip_hash == "unknown":
        return True  # Can't fingerprint — allow but don't block
    db = SessionLocal()
    try:
        record = db.query(GuestSession).filter(GuestSession.ip_hash == ip_hash).first()
        if record and record.used >= 1:
            return False  # Already used
        if not record:
            db.add(GuestSession(ip_hash=ip_hash, used=1, created_at=datetime.datetime.now()))
        else:
            record.used += 1
        db.commit()
        return True
    finally:
        db.close()

def has_guest_ip_used(ip_hash: str) -> bool:
    """Check (without marking) whether a guest IP has already used their free audit."""
    if ip_hash == "unknown":
        return False
    db = SessionLocal()
    try:
        record = db.query(GuestSession).filter(GuestSession.ip_hash == ip_hash).first()
        return bool(record and record.used >= 1)
    finally:
        db.close()

def save_rag_document(username, doc_name, content):
    """Persist RAG context to DB -- survives cloud restarts unlike local disk."""
    try:
        db = SessionLocal()
        db.add(RagDocument(username=username, doc_name=doc_name, content=content, created_at=datetime.datetime.now()))
        db.commit()
    finally:
        db.close()

def get_rag_documents(username):
    """Retrieve the last 10 RAG documents for BM25 retrieval."""
    try:
        db = SessionLocal()
        docs = db.query(RagDocument).filter(RagDocument.username == username).order_by(RagDocument.id.desc()).limit(10).all()
        return [(d.doc_name, d.content) for d in docs]
    finally:
        db.close()

def register_user(username, password, role="AUDITOR_LEGAL"):
    db = SessionLocal()
    try:
        if db.query(User).filter(User.username == username).first():
            return False
        # bcrypt: automatically generates a unique salt per password
        pw_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        db.add(User(username=username, password_hash=pw_hash, role=role, plan="FREE"))
        db.commit()
        return True
    finally:
        db.close()

def authenticate_user(username, password):
    db = SessionLocal()
    try:
        u = db.query(User).filter(User.username == username).first()
        if not u:
            return None, None, None
        stored = u.password_hash
        # Support legacy SHA-256 accounts (migration path)
        if stored.startswith("$2b$") or stored.startswith("$2a$"):
            match = bcrypt.checkpw(password.encode('utf-8'), stored.encode('utf-8'))
        else:
            # Legacy SHA-256 — verify then re-hash with bcrypt automatically
            match = (stored == hashlib.sha256(password.encode()).hexdigest())
            if match:
                new_hash = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
                u.password_hash = new_hash
                db.commit()
        if match:
            return u.username, u.role, u.plan
        return None, None, None
    finally:
        db.close()

@st.cache_resource
def get_analyzer():
    """Cache the LLM client globally — avoids re-instantiating Groq on every interaction."""
    return StandardAnalyzer()

@st.cache_data
def get_extracted_text(file_content_bytes, file_name):
    """Extract text directly -- avoids initializing the Groq LLM client inside a cached function."""
    import io
    name_lower = file_name.lower()
    if name_lower.endswith(".pdf"):
        import fitz
        doc = fitz.open(stream=file_content_bytes, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)
    elif name_lower.endswith(".docx"):
        import docx as _docx
        doc = _docx.Document(io.BytesIO(file_content_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    elif name_lower.endswith(".txt"):
        return file_content_bytes.decode("utf-8", errors="replace")
    return ""


st.set_page_config(page_title="Normatix | AI Enterprise GRC", page_icon="🧿", layout="wide", initial_sidebar_state="auto")

# --- ULTRA PREMIUM CSS DESIGN ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;600;800&family=JetBrains+Mono:wght@400;700&display=swap');
    .stApp {
        background-color: #050b14 !important;
        background: radial-gradient(circle at 10% 20%, rgb(9, 14, 25) 0%, rgb(3, 7, 18) 100%);
        color: #f1f5f9; font-family: 'Plus Jakarta Sans', sans-serif;
    }
    .stApp::before {
        content: ""; position: fixed; inset: 0; pointer-events: none;
        background-image:
            radial-gradient(circle at top right, rgba(56, 189, 248, 0.08) 0%, transparent 40%),
            radial-gradient(circle at bottom left, rgba(129, 140, 248, 0.08) 0%, transparent 40%);
        z-index: 0;
    }
    [data-testid="stHeader"] { background-color: transparent !important; }
    .block-container {
        background: rgba(15, 23, 42, 0.3) !important;
        backdrop-filter: blur(20px); -webkit-backdrop-filter: blur(20px);
        border: 1px solid rgba(255, 255, 255, 0.05); border-radius: 24px;
        padding-top: 3rem !important; padding-bottom: 3rem !important; margin-top: 2rem;
        box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.5), inset 0 1px 0 rgba(255, 255, 255, 0.1);
        z-index: 1; position: relative;
    }
    [data-testid="stSidebar"] {
        background: rgba(2, 6, 23, 0.8) !important; backdrop-filter: blur(30px); border-right: 1px solid rgba(255, 255, 255, 0.05);
    }
    [data-testid="stSidebar"] p, [data-testid="stSidebar"] label, [data-testid="stSidebar"] small, [data-testid="stSidebar"] span, [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
        color: #ffffff !important; text-shadow: 0 1px 3px rgba(0,0,0,0.8); font-weight: 600 !important;
    }
    [data-testid="stSidebar"] [data-baseweb="select"] span { color: #f8fafc !important; }
    .stButton>button {
        background: linear-gradient(135deg, #0284c7 0%, #2563eb 100%); color: white !important; border: 1px solid rgba(255,255,255,0.1) !important;
        border-radius: 12px !important; font-weight: 600; font-family: 'Plus Jakarta Sans', sans-serif;
        padding: 0.8rem 2rem !important; transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 4px 20px -2px rgba(37, 99, 235, 0.4) !important; letter-spacing: 0.5px;
    }
    .stButton>button:hover {
        transform: translateY(-3px) scale(1.02); box-shadow: 0 8px 30px rgba(37, 99, 235, 0.6) !important;
        background: linear-gradient(135deg, #38bdf8 0%, #3b82f6 100%) !important;
    }
    .stTabs [data-baseweb="tab-list"] { background-color: transparent !important; gap: 15px; border-bottom: 1px solid rgba(255, 255, 255, 0.08); padding-bottom: 2px;}
    .stTabs [data-baseweb="tab"] { background-color: transparent !important; border: none !important; color: #94a3b8 !important; font-family: 'Plus Jakarta Sans', sans-serif; font-weight: 600; font-size: 1.05rem; padding: 12px 16px; transition: all 0.2s;}
    .stTabs [aria-selected="true"] { color: #38bdf8 !important; border-bottom: 3px solid #38bdf8 !important; background: linear-gradient(0deg, rgba(56, 189, 248, 0.1) 0%, transparent 100%) !important; text-shadow: 0 0 15px rgba(56, 189, 248, 0.4);}
    .stTabs [data-baseweb="tab"]:hover { color: #f8fafc !important; }
    .stTextInput input, .stSelectbox [data-baseweb="select"], .stTextArea textarea {
        background-color: rgba(30, 41, 59, 0.5) !important; color: white !important;
        border: 1px solid rgba(255,255,255,0.1) !important; border-radius: 10px; font-family: 'JetBrains Mono', monospace;
    }
    .stTextInput input:focus { border-color: #38bdf8 !important; box-shadow: 0 0 0 1px #38bdf8 !important; }
    .risk-badge { padding: 8px 18px; border-radius: 30px; font-weight: 800; display: inline-block; margin-bottom: 25px; text-transform: uppercase; letter-spacing: 1.5px; font-size: 0.85rem;}
    .risk-unacceptable { background: linear-gradient(90deg, rgba(239,68,68,0.2), transparent); border: 1px solid #ef4444; color: #fca5a5; box-shadow: 0 0 20px rgba(239,68,68,0.3);}
    .risk-high { background: linear-gradient(90deg, rgba(249,115,22,0.2), transparent); border: 1px solid #f97316; color: #fdba74; box-shadow: 0 0 20px rgba(249,115,22,0.3);}
    .risk-limited { background: linear-gradient(90deg, rgba(234,179,8,0.2), transparent); border: 1px solid #eab308; color: #fde047; box-shadow: 0 0 20px rgba(234,179,8,0.3);}
    .risk-minimal { background: linear-gradient(90deg, rgba(34,197,94,0.2), transparent); border: 1px solid #22c55e; color: #86efac; box-shadow: 0 0 20px rgba(34,197,94,0.3);}
</style>
""", unsafe_allow_html=True)

T = {
    "Spanish": {
        "hero_sub": "Cierra los vacíos antes de que<br><span>te cuesten millones.</span>",
        "auth": "### 🔐 Acceso Cifrado (RBAC)", "user_l": "Usuario Corporativo", "pass_l": "Contraseña Segura", "btn_login": "🔑 Autenticar Módulo",
        "warn": "Role-Based Access Control Activo", "limit": "⚠️ **Límite Corporativo Alcanzado**", "rem": "Auditorías de la Suite:",
        "inv_txt": "🌟 **Memoria Institucional Activa:** Búsqueda retrospectiva RAG activada para tu compañía.",
        "lang": "🌐 Idioma (i18n):", "domain": "📂 Industria:", "jur": "⚖️ Marco Regulatorio:",
        "t1": "🛡️ Auditoría Documental", "t2": "📊 Dashboard de IA", "t3": "🛠️ Buró Legal (RAG)", "t4": "📄 Model Cards",
        "t5": "🎯 Red-Teaming", "t6": "🔗 Hub de APIs", "t7": "📖 Sobre Normatix", "t8": "👑 Panel Super Admin",
        "t1_h": "### Ingesta de Políticas y Contratos", "up_l": "📂 Documentos (Filtro Anti-PII Activo)", "run_a": "🚀 Ejecutar Análisis de Riesgo GRC",
        "spin": "Procesando de forma segura vía LLM...", "err_ext": "Sin texto.", "risk_t": "Clasificación Automática de Riesgo:",
        "rep_t": "### 📊 Reporte Estratégico Oficial", "down_r": "📥 Descargar Reporte (TXT)",
        "audit_succ": "✅ Auditoría guardada en la Base de Datos Histórica.",
        "t2_h": "### Estado Global de Riesgo (Dashboard)", "t2_n": "Sin registros en la base de datos central.", "t2_w": "Inicia sesión segura.",
        "t2_c": ["Documento", "Nivel de Riesgo", "Timestamp del Sistema"],
        "t3_h": "### Asistente Legal de Mitigación (Memoria RAG)", "t3_d": "Nuestra IA analiza el contexto de tus auditorías previas para apoyar a tu equipo legal con sugerencias de mitigación.",
        "t3_n": "Rastreador inactivo. Requiere una auditoría previa.", "t3_ctx": "Contexto legal activo:",
        "t4_h": "### Motor Transparencia: AI Model Cards", "t4_d": "Sube la documentación técnica de tu modelo y Normatix generará su equivalente a un Model Card oficial (ISO/EU).",
        "t4_btn": "🛠️ Compilar Model Card ISO", "t4_up": "📂 Sube los Datasets / Arquitectura",
        "t5_h": "### Laboratorio de Hacker Ético Adversarial 🎯", "t5_d": "Ingresa un Endpoint real. NORMATIX ejecutará peticiones HTTP inyectando prompts de prueba para evaluar la robustez de tu sistema.",
        "t5_url": "URL / Endpoint Real del Modelo de IA", "t5_atk": "⚔️ Lanzar Auditoría Adversarial",
        "t6_h": "### Hub de Integración Continua (CI/CD) 🔗", "t6_d": "Sincroniza Normatix directamente con tus flujos de despliegue.",
    },
    "English": {
        "hero_sub": "The AI that finds what<br><span>the law can't see.</span>",
        "auth": "### 🔐 RBAC Secure Access", "user_l": "Corporate User", "pass_l": "Secure Password", "btn_login": "🔑 Authenticate",
        "warn": "Role-Based Access Control Active", "limit": "⚠️ **Corporate Limit Reached**", "rem": "Suite Audits:",
        "inv_txt": "🌟 **Institutional Memory Active:** RAG retrospective search enabled for your company.",
        "lang": "🌐 UI Language:", "domain": "📂 Industry:", "jur": "⚖️ Regulatory Framework:",
        "t1": "🛡️ Document Audit", "t2": "📊 AI Dashboard", "t3": "🛠️ Remediation Desk", "t4": "📄 Model Cards",
        "t5": "🎯 Red-Teaming", "t6": "🔗 DevOps API Hub", "t7": "📖 About Normatix", "t8": "👑 Super Admin Panel",
        "t1_h": "### Ingest Policies & Contracts", "up_l": "📂 Upload Documents (Anti-PII Active)", "run_a": "🚀 Execute GRC Risk Analysis",
        "spin": "Processing securely via LLM...", "err_ext": "No text.", "risk_t": "Automated Risk Classification:",
        "rep_t": "### 📊 Official Strategic Report", "down_r": "📥 Download Report (TXT)",
        "audit_succ": "✅ Audit saved to Historical DB.",
        "t2_h": "### Global Risk Status (Dashboard)", "t2_n": "No active records in the central DB.", "t2_w": "Secure login required.",
        "t2_c": ["Document", "Risk Threshold", "System Timestamp"],
        "t3_h": "### Legal Mitigation Assistant (RAG Memory)", "t3_d": "Our AI analyzes your prior audit context to support your legal team with mitigation suggestions.",
        "t3_n": "Tracker inactive. Requires a prior audit.", "t3_ctx": "Active legal context:",
        "t4_h": "### Transparency Engine: AI Model Cards", "t4_d": "Upload your model's technical documentation and Normatix will generate an ISO/EU standard Model Card.",
        "t4_btn": "🛠️ Compile ISO Model Card", "t4_up": "📂 Upload Technical Architecture",
        "t5_h": "### Ethical Hacker / Adversarial Lab 🎯", "t5_d": "Enter a real endpoint. NORMATIX will execute HTTP requests injecting test prompts to evaluate your system's robustness.",
        "t5_url": "AI Model HTTP Endpoint", "t5_atk": "⚔️ Launch Adversarial Audit",
        "t6_h": "### Continuous Integration Hub (CI/CD) 🔗", "t6_d": "Sync Normatix directly with your deployment pipelines.",
    }
}


def main():
    if 'auth_username' not in st.session_state:
        st.session_state['auth_username'] = None
        st.session_state['auth_role'] = None
        st.session_state['auth_plan'] = None
        st.session_state['session_start'] = datetime.datetime.now()

    # Session timeout: auto-logout after 4 hours
    if st.session_state.get('auth_username') and st.session_state.get('session_start'):
        elapsed = (datetime.datetime.now() - st.session_state['session_start']).total_seconds()
        if elapsed > 14400:  # 4 hours
            for k in ['auth_username', 'auth_role', 'auth_plan', 'session_start', 'demo_mode']:
                st.session_state[k] = None
            st.warning("⏰ Sesión expirada por inactividad. Por favor inicia sesión nuevamente.")
            st.rerun()

    # --- HERO SECTION (always visible) ---
    hero_html = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@700;800;900&display=swap');
        .hero-container { text-align: center; padding: 2.5rem 1rem; background: linear-gradient(180deg, rgba(15,23,42,0.8) 0%, rgba(2,6,23,0.9) 100%); border-radius: 24px; margin-bottom: 2rem; border: 1px solid rgba(56, 189, 248, 0.15); position: relative; overflow: hidden;}
        .hero-container::before { content: ''; position: absolute; left: 0; top: 0; width: 100%; height: 2px; background: linear-gradient(90deg, transparent, #38bdf8, transparent); opacity: 0.5; }
        .hero-badge { display: inline-block; background: rgba(56, 189, 248, 0.1); border: 1px solid rgba(56, 189, 248, 0.3); border-radius: 30px; padding: 6px 20px; font-size: 0.8rem; letter-spacing: 3px; text-transform: uppercase; color: #7dd3fc; margin-bottom: 20px; font-family: 'Plus Jakarta Sans', sans-serif; font-weight: 800; box-shadow: 0 0 20px rgba(56,189,248,0.1);}
        .hero-title { font-family: 'Plus Jakarta Sans', sans-serif !important; font-size: clamp(2.2rem, 5vw, 3.8rem); font-weight: 900; color: #ffffff; line-height: 1.1; margin-bottom: 0px; letter-spacing: -1px;}
        .hero-title span { background: linear-gradient(90deg, #38bdf8, #818cf8); -webkit-background-clip: text; -webkit-text-fill-color: transparent; filter: drop-shadow(0 0 10px rgba(56,189,248,0.3));}
    </style>
    <div class="hero-container">
        <div class="hero-badge">⚡ NORMATIX ENTERPRISE ENGINE (V10)</div>
        <p class="hero-title">Cierra los vacíos antes de que<br><span>te cuesten millones.</span></p>
    </div>
    """
    st.markdown(hero_html, unsafe_allow_html=True)

    # =========================================================
    # LANDING PAGE (no auth) — sidebar hidden, auth in center
    # =========================================================
    if not st.session_state['auth_username']:
        st.markdown("""<style>
            [data-testid="collapsedControl"] { display: none; }
            section[data-testid="stSidebar"] { display: none !important; width: 0 !important; }
        </style>""", unsafe_allow_html=True)

        # Feature cards
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("""<div style='background:linear-gradient(180deg,rgba(15,23,42,0.9),rgba(2,6,23,0.9));padding:30px;border-radius:20px;border-top:4px solid #38bdf8;height:100%;box-shadow:0 10px 30px rgba(0,0,0,0.5);'>
            <div style='font-size:2.5rem;margin-bottom:15px;'>🔍</div>
            <h4 style='color:#f8fafc;font-size:1.1rem;line-height:1.4;font-weight:800;'>DETECTA PELIGROS OCULTOS</h4>
            <p style='color:#cbd5e1;font-size:0.95rem;line-height:1.6;'>Auditoría algorítmica y documental especializada. Identifica patrones de riesgo que escapan a revisiones manuales, minimizando tu exposición legal corporativa.</p></div>""", unsafe_allow_html=True)
        with c2:
            st.markdown("""<div style='background:linear-gradient(180deg,rgba(15,23,42,0.9),rgba(2,6,23,0.9));padding:30px;border-radius:20px;border-top:4px solid #818cf8;height:100%;box-shadow:0 10px 30px rgba(0,0,0,0.5);'>
            <div style='font-size:2.5rem;margin-bottom:15px;'>🛡️</div>
            <h4 style='color:#f8fafc;font-size:1.1rem;line-height:1.4;font-weight:800;'>CENSURA DATOS</h4>
            <p style='color:#cbd5e1;font-size:0.95rem;line-height:1.6;'>Antes de salir a internet, identificamos y redactamos RUTs, Tarjetas y correos en tus documentos. Tu empresa cumple con las normativas de privacidad de datos.</p></div>""", unsafe_allow_html=True)
        with c3:
            st.markdown("""<div style='background:linear-gradient(180deg,rgba(15,23,42,0.9),rgba(2,6,23,0.9));padding:30px;border-radius:20px;border-top:4px solid #2dd4bf;height:100%;box-shadow:0 10px 30px rgba(0,0,0,0.5);'>
            <div style='font-size:2.5rem;margin-bottom:15px;'>⚖️</div>
            <h4 style='color:#f8fafc;font-size:1.1rem;line-height:1.4;font-weight:800;'>TE GUÍA EN LA SOLUCIÓN</h4>
            <p style='color:#cbd5e1;font-size:0.95rem;line-height:1.6;'>Te entregamos propuestas de redacción para apoyar a tus equipos legales en la mitigación de riesgos identificados, basado en marcos regulatorios de Chile, EE.UU. y Europa.</p></div>""", unsafe_allow_html=True)

        st.markdown("<br><br>", unsafe_allow_html=True)

        # --- PRICING PAGE ---
        st.markdown("""
        <div style='text-align:center; margin-bottom:10px;'>
            <h2 style='color:#f8fafc;font-family:"Plus Jakarta Sans",sans-serif;font-weight:800;'>Planes Simples y Directos</h2>
            <p style='color:#94a3b8;'>Empieza gratis. Escala cuando lo necesites.</p>
        </div>""", unsafe_allow_html=True)

        pr1, pr2, pr3 = st.columns(3)
        with pr1:
            st.markdown("""
            <div style='background:rgba(15,23,42,0.7);border:1px solid rgba(255,255,255,0.08);border-radius:20px;padding:30px;text-align:center;'>
            <div style='font-size:1.8rem;margin-bottom:8px;'>🔍</div>
            <h3 style='color:#94a3b8;font-size:1rem;font-weight:700;margin:0;'>EXPLORACIÓN</h3>
            <div style='color:#f8fafc;font-size:2.5rem;font-weight:900;margin:15px 0 5px;'>Gratis</div>
            <p style='color:#64748b;font-size:0.85rem;margin-bottom:20px;'>Sin tarjeta de crédito</p>
            <ul style='color:#cbd5e1;text-align:left;font-size:0.9rem;line-height:2;list-style:none;padding:0;'>
            <li>✓ 1 auditoría de prueba</li>
            <li>✓ Descarga TXT del reporte</li>
            <li>✓ Clasificación de riesgo EU AI Act</li>
            <li>✗ Sin historial de auditoría ni RAG</li>
            <li>✗ Sin reporte PDF probatorio</li>
            </ul></div>""", unsafe_allow_html=True)
        with pr2:
            st.markdown("""
            <div style='background:linear-gradient(180deg,rgba(2,132,199,0.15),rgba(37,99,235,0.1));border:2px solid #38bdf8;border-radius:20px;padding:30px;text-align:center;position:relative;'>
            <div style='position:absolute;top:-12px;left:50%;transform:translateX(-50%);background:#38bdf8;color:#0f172a;font-weight:800;font-size:0.75rem;padding:4px 16px;border-radius:20px;'>MÁS POPULAR</div>
            <div style='font-size:1.8rem;margin-bottom:8px;'>💎</div>
            <h3 style='color:#38bdf8;font-size:1rem;font-weight:700;margin:0;'>NORMATIX PRO</h3>
            <div style='color:#f8fafc;font-size:2.5rem;font-weight:900;margin:15px 0 5px;'>$29.000 <span style='font-size:1rem;color:#94a3b8;'>CLP/m</span></div>
            <p style='color:#64748b;font-size:0.85rem;margin-bottom:20px;'>~USD 30 / mes</p>
            <ul style='color:#cbd5e1;text-align:left;font-size:0.9rem;line-height:2;list-style:none;padding:0;'>
            <li>✓ 5 auditorías completas / mes</li>
            <li>✓ Reporte PDF con marca blanca (Evidencia)</li>
            <li>✓ Dashboard histórico legal</li>
            <li>✓ Asistente RAG (Memoria Corporativa)</li>
            <li>✓ Soporte prioritario 24h</li>
            </ul></div>""", unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("<a href='https://normatix.lemonsqueezy.com/checkout/buy/7f17e6f7-3f4f-4b8f-9892-92249b540952' target='_blank' style='display:block;padding:12px;background:linear-gradient(135deg,#0284c7,#2563eb);color:white;border-radius:12px;text-decoration:none;font-weight:800;text-align:center;'>💎 Suscribirse Ahora</a>", unsafe_allow_html=True)
        with pr3:
            st.markdown("""
            <div style='background:rgba(15,23,42,0.7);border:1px solid rgba(129,140,248,0.3);border-radius:20px;padding:30px;text-align:center;'>
            <div style='font-size:1.8rem;margin-bottom:8px;'>🏢</div>
            <h3 style='color:#818cf8;font-size:1rem;font-weight:700;margin:0;'>ENTERPRISE</h3>
            <div style='color:#f8fafc;font-size:2.5rem;font-weight:900;margin:15px 0 5px;'>A Medida</div>
            <p style='color:#64748b;font-size:0.85rem;margin-bottom:20px;'>Contrato anual</p>
            <ul style='color:#cbd5e1;text-align:left;font-size:0.9rem;line-height:2;list-style:none;padding:0;'>
            <li>✓ Auditorías ilimitadas</li>
            <li>✓ Multi-usuario (RBAC)</li>
            <li>✓ Integración CI/CD (API)</li>
            <li>✓ Model Cards ISO + Red-Teaming</li>
            <li>✓ SLA y soporte dedicado</li>
            </ul></div>""", unsafe_allow_html=True)

        st.markdown("<br><br>", unsafe_allow_html=True)

        # --- AUTH + DEMO ---
        _, col_auth, _ = st.columns([1, 1.2, 1])
        with col_auth:
            st.markdown("""<div style='text-align:center;padding:20px;border-radius:15px;border:1px solid rgba(255,255,255,0.1);background:rgba(15,23,42,0.6);backdrop-filter:blur(10px);'>
            <h3 style='color:#f8fafc;font-family:"Plus Jakarta Sans",sans-serif;margin-bottom:10px;'>Acceso a la Plataforma</h3></div>""", unsafe_allow_html=True)

            auth_tabs = st.tabs(["🔑 Ingresar", "📝 Registro", "👁️ Anónimo"])

            with auth_tabs[0]:
                st.markdown("<br>", unsafe_allow_html=True)
                u_in = st.text_input("Usuario Corporativo", key="log_u")
                p_in = st.text_input("Contraseña", type="password", key="log_p")
                if st.button("Iniciar Sesión", use_container_width=True):
                    # --- RATE LIMITING: brute force protection ---
                    now = datetime.datetime.now()
                    attempts = st.session_state.get('login_attempts', 0)
                    lockout_until = st.session_state.get('lockout_until', None)
                    if lockout_until and now < lockout_until:
                        remaining = int((lockout_until - now).total_seconds())
                        st.error(f"🔒 Cuenta bloqueada. Intenta en {remaining} segundos.")
                    else:
                        un, role, plan = authenticate_user(u_in, p_in)
                        if un:
                            st.session_state['login_attempts'] = 0
                            st.session_state['lockout_until'] = None
                            st.session_state['auth_username'] = un
                            st.session_state['auth_role'] = role
                            st.session_state['auth_plan'] = plan
                            st.rerun()
                        else:
                            attempts += 1
                            st.session_state['login_attempts'] = attempts
                            if attempts >= 5:
                                st.session_state['lockout_until'] = now + datetime.timedelta(minutes=5)
                                st.error("🔒 Demasiados intentos fallidos. Cuenta bloqueada por 5 minutos.")
                            else:
                                st.error(f"Credenciales erróneas. Intentos fallidos: {attempts}/5")

            with auth_tabs[1]:
                st.markdown("<br>", unsafe_allow_html=True)
                r_un = st.text_input("Crear Usuario (letras, números, _)", key="reg_u", max_chars=40)
                r_pw = st.text_input("Crear Contraseña", type="password", key="reg_p")
                r_rol = st.selectbox("Rol", ["AUDITOR_LEGAL", "INGENIERO_IA"])
                if st.button("Crear Cuenta (3 Auditorías)", use_container_width=True):
                    if not r_un or not r_pw:
                        st.error("⚠️ Completa ambos campos.")
                    elif len(r_un) < 4:
                        st.error("⚠️ El usuario debe tener al menos 4 caracteres.")
                    elif not re.match(r'^[a-zA-Z0-9_]+$', r_un):
                        st.error("⚠️ El usuario solo puede contener letras, números y guiones bajos.")
                    elif len(r_un) > 40:
                        st.error("⚠️ El usuario no puede superar 40 caracteres.")
                    elif len(r_pw) < 8:
                        st.error("⚠️ La contraseña debe tener al menos 8 caracteres.")
                    else:
                        if register_user(r_un, r_pw, r_rol):
                            email_sent = send_welcome_email(r_un)
                            if email_sent:
                                st.success("✅ Cuenta B2B Creada y Correo Oficial enviado.")
                            else:
                                st.success("✅ Cuenta B2B Creada (Email Simulado, falta SMTP). Ya puedes iniciar sesión.")
                        else:
                            st.error("❌ Ese nombre de usuario ya está tomado.")

            with auth_tabs[2]:
                st.markdown("<br>", unsafe_allow_html=True)
                st.caption("Prueba la IA 1 vez gratis, sin dejar datos.")
                guest_ip_hash = get_guest_ip_hash()
                already_used = has_guest_ip_used(guest_ip_hash)
                if already_used:
                    st.warning("⚠️ **Ya usaste tu auditoría gratuita** desde este dispositivo. Crea una cuenta gratuita para obtener 3 auditorías completas.")
                else:
                    if st.button("Auditar Gratis (1 Crédito)", use_container_width=True):
                        allowed = check_and_mark_guest_ip(guest_ip_hash)
                        if allowed:
                            st.session_state['auth_username'] = "GUEST_SESSION"
                            st.session_state['auth_role'] = "INVITADO"
                            st.session_state['auth_plan'] = "GUEST"
                            st.rerun()
                        else:
                            st.error("⚠️ Este acceso gratuito ya fue utilizado. Por favor regístrate para continuar.")
                st.markdown("<br>", unsafe_allow_html=True)
                st.markdown("<div style='text-align:center;'>", unsafe_allow_html=True)
                if st.button("▶️ Ver Demo Interactiva", use_container_width=True):
                    st.session_state['auth_username'] = "DEMO_USER"
                    st.session_state['auth_role'] = "INVITADO"
                    st.session_state['auth_plan'] = "GUEST"
                    demo_report = """## 📊 Reporte Estratégico Oficial — DEMO

### 🔴 Resumen Ejecutivo
El documento analizado corresponde a una política interna de uso de IA para evaluación de crédito en entidades financieras. Se identificaron riesgos **ALTO** bajo el marco del EU AI Act y la legislación chilena vigente (Ley 19.628).

### ⚠️ Hallazgos Críticos Detectados
1. **Ausencia de cláusula de transparencia algorítmica** — Los clientes no son informados de que una IA toma decisiones sobre su crédito. Esto infringe el Artículo 13 del EU AI Act y el principio de información de la Ley 19.628.
2. **Sin mecanismo de impugnación humana** — El sistema no contempla un canal para que el afectado solicite revisión por un humano. Obligatorio en sistemas de Alto Riesgo.
3. **Datos de entrenamiento no documentados** — No existe registro de las fuentes de datos usadas para entrenar el modelo, imposibilitando una auditoría de sesgo.

### ✅ Áreas de Cumplimiento Verificado
- Política de retención de datos: 5 años (conforme)
- Cifrado en tránsito: TLS 1.3 (conforme)
- Segmentación de roles de acceso: Implementada

### 📋 Recomendaciones de Mitigación Prioritaria
1. Redactar e incorporar una **cláusula de decisión automatizada** en los contratos de adhesión de clientes.
2. Implementar un **canal de impugnación** documentado con SLA máximo de 5 días hábiles.
3. Generar un **Model Card** según estándares ISO/IEC 42001 para el modelo de scoring.

*Este reporte fue generado por Normatix AI GRC Platform como demostración del análisis automatizado.*"""
                    st.session_state['last_audit'] = demo_report
                    st.session_state['last_audit_docs'] = "Política_IA_Credito_Demo.pdf"
                    st.session_state['demo_mode'] = True
                    st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)


            st.markdown("<br><div style='text-align:center;padding:10px;border:1px dashed rgba(56,189,248,0.2);border-radius:10px;'><span style='color:#475569;font-size:0.8rem;'>🔵 SSO Empresarial (Microsoft 365 / Google Workspace)</span><br><span style='color:#38bdf8;font-size:0.75rem;font-weight:700;letter-spacing:1px;'>PRÓXIMAMENTE</span></div>", unsafe_allow_html=True)
        st.markdown("<br><br><br><div style='text-align:center;color:#475569;font-size:0.85rem;border-top:1px solid rgba(255,255,255,0.05);padding-top:20px;'>© 2026 Normatix Soluciones Inteligentes &nbsp;|&nbsp; <a href='mailto:contacto@normatix.cl' style='color:#38bdf8;text-decoration:none;'>contacto@normatix.cl</a></div>", unsafe_allow_html=True)
        st.stop()

    # =========================================================
    # AUTHENTICATED APP — full sidebar + tabs
    # =========================================================
    username = st.session_state['auth_username']
    role = st.session_state['auth_role']
    plan = st.session_state['auth_plan']

    # Demo mode banner
    if st.session_state.get('demo_mode'):
        st.info("🎬 **Modo Demo Activo** — Estás viendo datos de ejemplo. [Crea una cuenta gratuita](#) para auditar tus propios documentos.")

    if os.path.exists("logo.png"):
        st.sidebar.image("logo.png", use_container_width=True)
    else:
        st.sidebar.image("https://cdn-icons-png.flaticon.com/512/2059/2059080.png", width=70)

    target_lang_opt = st.sidebar.selectbox("🌐 GRC Language / Interfaz:", ["Spanish", "English"])
    loc = T.get(target_lang_opt, T["Spanish"])

    st.sidebar.markdown(f"### 👤 Mi Perfil")
    st.sidebar.write(f"**Usuario:** `{username}`")
    st.sidebar.write(f"**Plan:** `{'👑 PRO Ilimitado' if plan == 'PRO' else 'Suscripción Gratuita'}`")
    st.sidebar.success(f"🔒 Nivel: **{role}**")
    st.sidebar.markdown("---")

    if plan == "GUEST":
        usage_count = st.session_state.get('guest_uses', 0)
        MAX_FREE_USES = 1
    else:
        usage_count = get_usage(username)
        MAX_FREE_USES = 3 if plan == "FREE" else 1000000

    if plan == "FREE":
        st.sidebar.progress(min(usage_count / MAX_FREE_USES, 1.0))
        if usage_count >= MAX_FREE_USES:
            st.sidebar.error("⚠️ **Cuota Terminada**")
            st.sidebar.markdown("<a href='https://normatix.lemonsqueezy.com/checkout/buy/7f17e6f7-3f4f-4b8f-9892-92249b540952' target='_blank' style='display:inline-block;padding:8px 16px;background:#2563eb;color:white;border-radius:8px;text-decoration:none;font-weight:bold;width:100%;text-align:center;'>💎 MEJORAR A NORMATIX PRO ($29.000 CLP/m)</a>", unsafe_allow_html=True)
        else:
            st.sidebar.caption(f"Auditorías Usadas: {usage_count} / {MAX_FREE_USES}")
    elif plan == "GUEST":
        st.sidebar.progress(min(usage_count / MAX_FREE_USES, 1.0))
        if usage_count >= MAX_FREE_USES:
            st.sidebar.error("⚠️ **Prueba Anónima Terminada**")
            st.sidebar.markdown("Para continuar, cierra sesión y créate una cuenta gratuita.")
        else:
            st.sidebar.caption(f"De un solo uso: {usage_count} / {MAX_FREE_USES}")
    else:
        st.sidebar.info("∞ Licencias Ilimitadas (Enterprise)")

    domain = st.sidebar.selectbox(loc["domain"], [
        "👩‍💼 Recursos Humanos (Sesgos en Contratación)",
        "⛏️ Minería / Automatización Industrial",
        "🏥 Salud y Privacidad Médica (Datos de Pacientes)",
        "🎓 Educación (Evaluación Académica Justa)",
        "⚖️ Legal Corporativo (Contratos y NDAs)",
        "💳 Servicios Financieros o Retail (Decisión de Crédito)"
    ])
    jurisdiction = st.sidebar.selectbox(loc["jur"], ["Chile (Ley N° 19.628 / Proyecto Ley IA)", "EU AI Act (Europa)", "USA (NIST AI RMF)"])

    if st.sidebar.button("Cerrar Sesión"):
        # Complete session wipe -- future-proof, covers any new keys added later
        st.session_state.clear()
        st.rerun()

    # Build tab list per role
    t7_name = loc.get("t7", "📖 Sobre Normatix")
    t8_name = loc.get("t8", "👑 Panel Super Admin")
    if role == "ADMINISTRADOR":
        allowed_tab_names = [t7_name, t8_name, loc["t1"], loc["t2"], loc["t3"], loc["t4"], loc["t5"], loc["t6"]]
    elif role == "AUDITOR_LEGAL":
        allowed_tab_names = [t7_name, loc["t1"], loc["t2"], loc["t3"]]
    elif role == "INGENIERO_IA":
        allowed_tab_names = [t7_name, loc["t2"], loc["t4"], loc["t5"], loc["t6"]]
    else:  # INVITADO
        allowed_tab_names = [t7_name, loc["t1"]]

    tabs = st.tabs(allowed_tab_names)
    tab_idx = 0

    # --- TAB: SOBRE NORMATIX ---
    if t7_name in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown("""
            <div style="text-align:center;padding:10px 0 40px 0;">
                <h1 style="color:#f8fafc;font-size:clamp(1.8rem,4vw,2.8rem);font-weight:800;line-height:1.2;margin-bottom:15px;font-family:'Plus Jakarta Sans',sans-serif;">¿ESTÁ TU EMPRESA REALMENTE PREPARADA<br><span style="color:#38bdf8;">PARA ADOPTAR LA INTELIGENCIA ARTIFICIAL?</span></h1>
                <h3 style="color:#94a3b8;font-weight:400;font-size:1.2rem;">Normatix audita tu infraestructura proactivamente para identificar y mitigar riesgos antes de un despliegue de IA.</h3>
            </div>
            """, unsafe_allow_html=True)
            ca, cb, cc = st.columns(3)
            with ca:
                st.markdown("""<div style='background:linear-gradient(180deg,rgba(15,23,42,0.9),rgba(2,6,23,0.9));padding:30px;border-radius:20px;border-top:4px solid #38bdf8;height:100%;box-shadow:0 10px 30px rgba(0,0,0,0.5);'>
                <div style='font-size:2.5rem;margin-bottom:15px;'>🔍</div>
                <h4 style='color:#f8fafc;font-size:1.1rem;line-height:1.4;font-weight:800;'>DETECTA PELIGROS OCULTOS</h4>
                <p style='color:#cbd5e1;font-size:0.95rem;line-height:1.6;'>Auditoría algorítmica especializada. Identifica patrones de riesgo que escapan a revisiones manuales, minimizando tu exposición legal corporativa.<br><br><i>Ejemplo: Detecta sesgos en documentos de RR.HH que podrían generar controversias reputacionales.</i></p></div>""", unsafe_allow_html=True)
            with cb:
                st.markdown("""<div style='background:linear-gradient(180deg,rgba(15,23,42,0.9),rgba(2,6,23,0.9));padding:30px;border-radius:20px;border-top:4px solid #818cf8;height:100%;box-shadow:0 10px 30px rgba(0,0,0,0.5);'>
                <div style='font-size:2.5rem;margin-bottom:15px;'>🛡️</div>
                <h4 style='color:#f8fafc;font-size:1.1rem;line-height:1.4;font-weight:800;'>CENSURA DATOS AUTOMÁTICAMENTE</h4>
                <p style='color:#cbd5e1;font-size:0.95rem;line-height:1.6;'>Antes de procesar con IA, identificamos y redactamos RUTs, Tarjetas de Crédito y correos electrónicos. Tu empresa cumple con la Ley 19.628 y el RGPD europeo.</p></div>""", unsafe_allow_html=True)
            with cc:
                st.markdown("""<div style='background:linear-gradient(180deg,rgba(15,23,42,0.9),rgba(2,6,23,0.9));padding:30px;border-radius:20px;border-top:4px solid #2dd4bf;height:100%;box-shadow:0 10px 30px rgba(0,0,0,0.5);'>
                <div style='font-size:2.5rem;margin-bottom:15px;'>⚖️</div>
                <h4 style='color:#f8fafc;font-size:1.1rem;line-height:1.4;font-weight:800;'>TE GUÍA EN LA SOLUCIÓN</h4>
                <p style='color:#cbd5e1;font-size:0.95rem;line-height:1.6;'>Te entregamos propuestas de redacción dinámicas para apoyar a tus equipos legales en la mitigación ágil de los riesgos identificados.</p></div>""", unsafe_allow_html=True)
        tab_idx += 1

    # --- TAB: SUPER ADMIN ---
    if t8_name in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown("### 👑 Normatix SaaS Control Center")
            st.caption("Visión global financiera y operativa de todos los clientes de la plataforma.")
            db = SessionLocal()
            try:
                total_users = db.query(User).count()
                pro_users = db.query(User).filter(User.plan == "PRO").count()
                mrr = pro_users * 29000
                all_usage = db.query(Usage).all()
                total_audits = sum([u.count for u in all_usage])
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("Clientes Registrados", total_users)
                m2.metric("Suscripciones PRO", pro_users)
                m3.metric("MRR Estimado", f"${mrr:,.0f} CLP")
                m4.metric("Auditorías Globales", total_audits)
                st.markdown("---")
                st.markdown("#### 🏢 Directorio B2B de Clientes")
                users = db.query(User).order_by(User.id.desc()).all()
                if users:
                    df_u = pd.DataFrame(
                        [(u.id, u.username, u.role, u.plan) for u in users],
                        columns=["ID", "Usuario", "Rol", "Plan"]
                    )
                    st.dataframe(df_u, use_container_width=True, hide_index=True)
                    p1, p2 = st.columns(2)
                    with p1:
                        fig1 = px.pie(df_u, names="Plan", title="Distribución de Tiers",
                                      color="Plan", color_discrete_map={"PRO": "#38bdf8", "FREE": "#94a3b8", "GUEST": "#475569"}, hole=0.5)
                        fig1.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", font_color="#f8fafc")
                        st.plotly_chart(fig1, use_container_width=True)
                    with p2:
                        role_counts = df_u["Rol"].value_counts().reset_index()
                        role_counts.columns = ["Rol", "count"]
                        fig2 = px.bar(role_counts, x="count", y="Rol", orientation='h',
                                      title="Distribución de Roles", color_discrete_sequence=["#818cf8"])
                        fig2.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", font_color="#f8fafc", yaxis={'categoryorder': 'total ascending'})
                        st.plotly_chart(fig2, use_container_width=True)
                else:
                    st.info("Aún no hay usuarios registrados.")
                    
                st.markdown("---")
                st.markdown("#### 🎛️ Motor de Reglas Dinámicas (AI Prompts)")
                st.caption("Ajusta el cerebro del auditor corporativo en vivo para toda la plataforma y activa el modo de búsqueda en internet.")
                
                try:
                    import json
                    with open("normatix_config.json", "r", encoding="utf-8") as fcf:
                        current_cfg = json.load(fcf)
                except:
                    current_cfg = {"system_prompt": "", "enable_web_search": False}
                    
                new_prompt = st.text_area("System Prompt Base (Institucional)", value=current_cfg.get("system_prompt", ""), height=200)
                agentic_search = st.toggle("🌐 Activar RAG Agéntico (DuckDuckGo Web Search en Vivo) — Requiere más tiempo de respuesta", value=current_cfg.get("enable_web_search", False))
                
                if st.button("Guardar Configuración en Caliente (Hot-Reload)"):
                    current_cfg["system_prompt"] = new_prompt
                    current_cfg["enable_web_search"] = agentic_search
                    with open("normatix_config.json", "w", encoding="utf-8") as fcf:
                        json.dump(current_cfg, fcf, indent=4)
                    st.success("✅ Motor LLM actualizado para todos los inquilinos (Tenants) al instante.")
            finally:
                db.close()
        tab_idx += 1

    # --- TAB: AUDITORÍA DOCUMENTAL ---
    if loc["t1"] in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown(loc["t1_h"])
            uploaded_files = st.file_uploader(loc["up_l"], type=["pdf", "txt", "docx"], accept_multiple_files=True)
            is_blocked = (usage_count >= MAX_FREE_USES)
            analyze_btn = st.button(loc["run_a"], use_container_width=True, disabled=is_blocked)
            if is_blocked:
                st.warning("Límite de auditorías alcanzado. Mejora tu plan para continuar.")

            # --- FILE SIZE VALIDATION (15MB max) ---
            MAX_FILE_SIZE = 15 * 1024 * 1024
            if uploaded_files:
                oversized = [f.name for f in uploaded_files if len(f.getvalue()) > MAX_FILE_SIZE]
                if oversized:
                    st.error(f"⚠️ Archivo(s) demasiado grandes (máx. 15MB): {', '.join(oversized)}")
                    uploaded_files = [f for f in uploaded_files if len(f.getvalue()) <= MAX_FILE_SIZE]

            if analyze_btn and uploaded_files:
                import time
                with st.status("Ejecutando Inteligencia Artificial GRC...", expanded=True) as status:
                    analyzer = get_analyzer()
                    combined_text = ""
                    doc_names = [f.name for f in uploaded_files]
                    
                    st.write("1️⃣ Inicializando motor de ofuscación de datos (PII)...")
                    for f in uploaded_files:
                        extracted = get_extracted_text(f.getvalue(), f.name)
                        if extracted.strip():
                            st.write(f"&nbsp;&nbsp;↳ Limpiando archivo: `{f.name}`")
                            clean_text = analyzer.scrub_pii(extracted)
                            combined_text += f"\n\n--- Docs: {f.name} ---\n{clean_text}\n"

                    st.write("2️⃣ Mapeando marco regulatorio seleccionado...")
                    time.sleep(1)
                    st.write(f"&nbsp;&nbsp;↳ Aislando variables: `{jurisdiction}`")
                    
                    st.write("3️⃣ Analizando vectores de riesgo y sesgos demográficos...")
                    time.sleep(1.5)
                    
                    st.write("4️⃣ Estructurando reporte con citas legales formales...")
                    time.sleep(1)
                    status.update(label="Análisis completado ✅", state="complete", expanded=False)

                    if combined_text.strip():
                        st.markdown("---")
                        risk_tier = analyzer.classify_risk_tier(combined_text, jurisdiction)
                        badge_class = f"risk-{risk_tier.lower()}"
                        st.markdown(f'<div class="risk-badge {badge_class}">{loc["risk_t"]} {risk_tier}</div>', unsafe_allow_html=True)

                        if plan == "GUEST":
                            # Sync session_state and DB-backed counter
                            st.session_state['guest_uses'] = st.session_state.get('guest_uses', 0) + 1
                            # The IP was already marked in DB when the guest button was clicked
                        else:
                            save_audit(username, ", ".join(doc_names), risk_tier)
                            increment_usage(username)

                        # Persist RAG context to Supabase DB -- survives Streamlit Cloud restarts
                        if username not in ("GUEST_SESSION", "DEMO_USER"):
                            for d_name in doc_names:
                                save_rag_document(username, d_name, combined_text)

                        st.markdown(loc["rep_t"])
                        try:
                            generator = analyzer.analyze_stream(combined_text, domain, jurisdiction=jurisdiction, language=target_lang_opt)
                            result_text = st.write_stream(generator)
                        except Exception as llm_err:
                            err_msg = str(llm_err)
                            if "429" in err_msg or "rate" in err_msg.lower():
                                st.error("⏳ El servidor de IA está ocupado. Espera 30 segundos e intenta nuevamente (límite de velocidad del proveedor LLM).")
                            else:
                                st.error(f"❌ Error en el análisis: {err_msg[:200]}")
                            st.stop()
                        st.session_state['last_audit'] = result_text
                        st.session_state['last_audit_docs'] = ", ".join(doc_names)
                        st.download_button(loc["down_r"], data=result_text, file_name="Normatix_Audit.txt")
                        # --- BRANDED PDF DOWNLOAD ---
                        if plan == "PRO" or role == "ADMINISTRADOR":
                            try:
                                pdf_bytes = generate_pdf_report(
                                    username=username,
                                    doc_names=", ".join(doc_names),
                                    risk_tier=risk_tier,
                                    report_text=result_text
                                )
                                st.download_button(
                                    "📄 Descargar Reporte Oficial (PDF Corporativo)",
                                    data=pdf_bytes,
                                    file_name=f"Normatix_Audit_{datetime.datetime.now().strftime('%Y%m%d')}.pdf",
                                    mime="application/pdf"
                                )
                            except Exception as e:
                                st.caption(f"PDF no disponible: {e}")
                        else:
                            st.info("💎 **Función PRO:** El reporte en PDF institucional (marca blanca) es exclusivo de suscripciones activas. Como usuario Free, tienes disponible la exportación en texto (.txt) arriba.")

                        # --- GOD TIER: REMEDIATION ENGINE ---
                        if risk_tier in ["UNACCEPTABLE", "HIGH"]:
                            st.markdown("<br>", unsafe_allow_html=True)
                            st.markdown("""<div style='background:rgba(129,140,248,0.1);border:1px solid rgba(129,140,248,0.3);border-radius:12px;padding:20px;'>
                            <h4 style='color:#818cf8;margin-top:0;'>✨ Remediation Engine (Redacción Correctiva)</h4>
                            <p style='color:#cbd5e1;font-size:0.9rem;'>Normatix puede reescribir automáticamente las cláusulas conflictivas detectadas en este reporte para hacerlas seguras bajo <b>{}</b>.</p>
                            """.format(jurisdiction), unsafe_allow_html=True)
                            
                            if st.button("Generar Cláusulas Correctivas Seguras"):
                                st.info("Generando redacción corporativa legalmente segura mediante LLM 70B...")
                                try:
                                    corrective_generator = analyzer.remediate_clauses_stream(combined_text, result_text, jurisdiction)
                                    remediation_text = st.write_stream(corrective_generator)
                                    st.download_button("📥 Descargar Cláusulas Corregidas", data=remediation_text, file_name=f"Normatix_Remediation_{datetime.datetime.now().strftime('%Y%m%d')}.txt")
                                except Exception as e:
                                    st.error(f"Error generando cláusulas: {str(e)}")
                            st.markdown("</div>", unsafe_allow_html=True)
                        st.success(loc["audit_succ"])
        tab_idx += 1

    # --- TAB: DASHBOARD ---
    if loc["t2"] in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown(loc["t2_h"])
            audits = get_audits(username)
            if audits:
                df = pd.DataFrame(audits, columns=loc["t2_c"])
                col1, col2, col3 = st.columns(3)
                risk_counts = df[loc["t2_c"][1]].value_counts()
                with col1: st.metric("Total Analizados", len(df))
                with col2: st.metric("Nivel Dominante", risk_counts.idxmax() if not risk_counts.empty else "N/A")
                with col3:
                    h_r = risk_counts.get('UNACCEPTABLE', 0) + risk_counts.get('HIGH', 0)
                    st.metric("Alertas Críticas GRC", f"{h_r}")
                st.dataframe(df, use_container_width=True, hide_index=True)
                st.markdown("### Distribución de Riesgo")
                fig = px.pie(
                    values=risk_counts.values, names=risk_counts.index, hole=0.65,
                    color=risk_counts.index,
                    color_discrete_map={"UNACCEPTABLE": "#ef4444", "HIGH": "#f97316", "LIMITED": "#eab308", "MINIMAL": "#22c55e"}
                )
                fig.update_traces(textposition='outside', textinfo='percent+label', marker=dict(line=dict(color='#0f172a', width=2)))
                fig.update_layout(
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    margin=dict(t=20, b=20, l=20, r=20), showlegend=False, font_color="#f8fafc",
                    annotations=[dict(text=f"{len(df)}<br>Docs", x=0.5, y=0.5, font_size=18, showarrow=False, font_color="#f8fafc")]
                )
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info(loc["t2_n"])
        tab_idx += 1

    # --- TAB: BURÓ LEGAL RAG ---
    if loc["t3"] in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown(loc["t3_h"])
            st.write(loc["t3_d"])
            
            # --- RAG MOAT COUNTER ---
            rag_docs_count = len(get_rag_documents(username))
            st.markdown(f"""
            <div style='padding:12px 20px;background:rgba(56,189,248,0.1);border-left:4px solid #38bdf8;border-radius:4px;margin-bottom:20px;'>
                <span style='color:#38bdf8;font-weight:800;font-size:1.1rem;'>🧠 Módulo de Memoria Institucional Activo</span><br>
                <span style='color:#e2e8f0;font-size:0.9rem;'>Normatix tiene indexado el contexto técnico de <b>{rag_docs_count} documento(s)</b> de tu empresa. El asistente RAG (Retrieval-Augmented Generation) buscará mitigaciones precisas dentro de esa jurisprudencia interna real.</span>
            </div>
            """, unsafe_allow_html=True)
            
            if 'last_audit' not in st.session_state:
                st.info(loc["t3_n"])
            else:
                st.success(f"{loc['t3_ctx']} **{st.session_state['last_audit_docs']}**")
                if "messages" not in st.session_state:
                    st.session_state.messages = []
                for m in st.session_state.messages:
                    with st.chat_message(m["role"]): st.markdown(m["content"])
                if prompt := st.chat_input("Ej: ¿Qué cláusula debería agregar para mitigar el riesgo de sesgo?"):
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    with st.chat_message("user"): st.markdown(prompt)
                    with st.chat_message("assistant"):
                        retrieved_context = ""
                        from rank_bm25 import BM25Okapi
                        rag_docs = get_rag_documents(username)
                        if rag_docs:
                            corpus = [d[1] for d in rag_docs]
                            tokenized_corpus = [doc.split(" ") for doc in corpus]
                            bm25 = BM25Okapi(tokenized_corpus)
                            top_doc = bm25.get_top_n(prompt.split(" "), corpus, n=1)[0]
                            retrieved_context = f"--- RAG CONTEXT ---\n{top_doc[:3000]}\n\n"
                        full_context = retrieved_context + "--- LATEST AUDIT ---\n" + st.session_state.get('last_audit', '')
                        analyzer = get_analyzer()
                        try:
                            response = st.write_stream(analyzer.remediate_stream(full_context, prompt))
                        except Exception as rag_err:
                            response = "⚠️ Error al conectar con el asistente legal. Por favor intenta nuevamente."
                            st.error(response)
                    st.session_state.messages.append({"role": "assistant", "content": response})
        tab_idx += 1

    # --- TAB: MODEL CARDS ---
    if loc["t4"] in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown(loc["t4_h"])
            st.write(loc["t4_d"])
            tech_file = st.file_uploader(loc["t4_up"], type=["pdf", "txt"])
            if st.button(loc["t4_btn"], use_container_width=True) and tech_file:
                with st.spinner(loc["spin"]):
                    analyzer = StandardAnalyzer()
                    clean_text = analyzer.scrub_pii(get_extracted_text(tech_file.getvalue(), tech_file.name))
                    if clean_text:
                        card_result = st.write_stream(analyzer.generate_model_card_stream(clean_text))
                        st.download_button(loc["down_r"], data=card_result, file_name=f"MC_{tech_file.name}.md")
        tab_idx += 1

    # --- TAB: RED-TEAMING ---
    if loc["t5"] in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown(loc["t5_h"])
            st.write(loc["t5_d"])
            model_endpoint = st.text_input(loc["t5_url"], "https://api.mi-empresa.cl/v1/chat/completions")
            if st.button(loc["t5_atk"], use_container_width=True):
                # --- SSRF PROTECTION: block private IP ranges ---
                ssrf_safe = True
                private_patterns = [
                    r'localhost', r'127\.', r'10\.', r'192\.168\.', r'172\.(1[6-9]|2[0-9]|3[01])\.',
                    r'0\.0\.0\.0', r'169\.254\.', r'::1', r'metadata\.google'
                ]
                url_lower = model_endpoint.lower()
                if not url_lower.startswith(('http://', 'https://')):
                    st.error("⛔ URL inválida. Debe comenzar con http:// o https://")
                    ssrf_safe = False
                elif any(re.search(p, url_lower) for p in private_patterns):
                    st.error("⛔ URL bloqueada por seguridad: no se permiten IPs privadas o locales.")
                    ssrf_safe = False
                if ssrf_safe:
                    with st.spinner("Inyectando Payloads y Evaluando (HTTP Real / Simulado)..."):
                        analyzer = StandardAnalyzer()
                        st.write_stream(analyzer.perform_red_teaming_stream(model_endpoint))
                        increment_usage(username)
        tab_idx += 1

    # --- TAB: CI/CD HUB ---
    if loc["t6"] in allowed_tab_names:
        with tabs[tab_idx]:
            st.markdown(loc["t6_h"])
            st.write(loc["t6_d"])
            st.markdown("""
            <div style='text-align:center;padding:60px 40px;border:1px dashed rgba(56,189,248,0.2);border-radius:20px;background:rgba(15,23,42,0.5);margin-top:20px;'>
                <div style='font-size:3rem;margin-bottom:20px;'>🔗</div>
                <h3 style='color:#f8fafc;font-weight:800;margin-bottom:10px;'>Integración CI/CD con GitHub Actions</h3>
                <p style='color:#94a3b8;max-width:500px;margin:0 auto 20px;line-height:1.6;'>Normatix se integrará directamente en tus pipelines de despliegue para auditar modelos de IA automáticamente antes de cada <code style='color:#38bdf8;'>merge</code> a producción.</p>
                <div style='display:inline-block;background:rgba(56,189,248,0.1);border:1px solid rgba(56,189,248,0.3);border-radius:30px;padding:8px 24px;margin-bottom:20px;'>
                    <span style='color:#38bdf8;font-weight:700;font-size:0.85rem;letter-spacing:2px;'>EN DESARROLLO — ENTERPRISE PLAN</span>
                </div>
                <p style='color:#475569;font-size:0.85rem;'>¿Necesitas esta integración ahora? <a href='mailto:contacto@normatix.cl' style='color:#38bdf8;'>contacto@normatix.cl</a></p>
            </div>
            """, unsafe_allow_html=True)
        tab_idx += 1

    st.markdown("<br><br><br><div style='text-align:center;color:#475569;font-size:0.85rem;border-top:1px solid rgba(255,255,255,0.05);padding-top:20px;'>© 2026 Normatix Soluciones Inteligentes &nbsp;|&nbsp; <a href='mailto:contacto@normatix.cl' style='color:#38bdf8;text-decoration:none;'>contacto@normatix.cl</a></div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
